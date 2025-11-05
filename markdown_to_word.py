"""
Markdown to Word Document Converter with LaTeX to OMML Equation Support

This module provides functionality to convert markdown files to Word documents,
automatically converting embedded LaTeX equations to OMML (Office Math Markup Language)
format for proper rendering in Microsoft Word.
"""

import subprocess
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from pathlib import Path


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def latex_to_omml(latex_formula):
    """
    Convert LaTeX formula to OMML (Office Math Markup Language) using texmath.
    
    Args:
        latex_formula (str): LaTeX string (without $ or $$ delimiters)
    
    Returns:
        str or None: OMML string or None if conversion fails
    """
    try:
        # Use full path to texmath executable
        texmath_path = os.path.expanduser("~/.local/bin/texmath")
        
        # Call texmath to convert LaTeX to OMML
        result = subprocess.run(
            [texmath_path, '--from', 'tex', '--to', 'omml'],
            input=latex_formula.encode('utf-8'),
            capture_output=True,
            timeout=5
        )
        
        if result.returncode == 0:
            return result.stdout.decode('utf-8').strip()
        else:
            print(f"Error converting formula: {result.stderr.decode('utf-8')}")
            return None
    except Exception as e:
        print(f"Exception during conversion: {e}")
        return None


def extract_latex_equations(content):
    """
    Extract LaTeX equations from markdown content.
    
    Extracts both display equations ($$...$$) and inline equations ($...$).
    
    Args:
        content (str): Markdown content as string
    
    Returns:
        list: List of tuples (equation, is_display_mode) where is_display_mode is bool
    """
    equations = []
    
    # First extract display equations ($$...$$)
    display_pattern = r'\$\$(.*?)\$\$'
    for match in re.finditer(display_pattern, content, re.DOTALL):
        eq = match.group(1).strip()
        equations.append((eq, True))  # True = display mode
    
    # Remove display equations from content for inline extraction
    temp_content = re.sub(display_pattern, '', content, flags=re.DOTALL)
    
    # Then extract inline equations ($...$)
    inline_pattern = r'\$([^\$]+?)\$'
    for match in re.finditer(inline_pattern, temp_content):
        eq = match.group(1).strip()
        if eq and '\n' not in eq:  # Skip if multiline
            equations.append((eq, False))  # False = inline mode
    
    return equations


def convert_equations_to_omml(equations_list, verbose=False):
    """
    Convert a list of LaTeX equations to OMML format.
    
    Args:
        equations_list (list): List of tuples (equation, is_display)
        verbose (bool): If True, print progress information
    
    Returns:
        list: List of tuples (omml, is_display) for successfully converted equations
    """
    omml_equations = []
    
    if verbose:
        print("Converting extracted equations to OMML:\n")
    
    for i, (latex_eq, is_display) in enumerate(equations_list, 1):
        mode = "DISPLAY" if is_display else "INLINE"
        if verbose:
            print(f"Converting equation {i} [{mode}]: {latex_eq[:40]}...")
        
        omml = latex_to_omml(latex_eq)
        if omml:
            omml_equations.append((omml, is_display))
            if verbose:
                print(f"  ✓ Success")
        else:
            if verbose:
                print(f"  ✗ Failed")
    
    if verbose:
        print(f"\nSuccessfully converted {len(omml_equations)} out of {len(equations_list)} equations")
    
    return omml_equations


# ============================================================================
# DOCUMENT GENERATION
# ============================================================================

def create_word_doc_from_markdown(markdown_content, omml_equations_data, output_path="output.docx", verbose=False):
    """
    Create a Word document from markdown content with OMML equations.
    
    Args:
        markdown_content (str): Full markdown text
        omml_equations_data (list): List of (omml, is_display) tuples
        output_path (str): Path where the document will be saved
        verbose (bool): If True, print progress information
    
    Returns:
        str: Path to the created document
    """
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    # Create a new Document
    doc = Document()
    
    # Create a mapping of equation positions
    display_eqs = [omml for omml, is_display in omml_equations_data if is_display]
    inline_eqs = [omml for omml, is_display in omml_equations_data if not is_display]
    
    # Index to track which OMML equation we're on
    display_eq_index = 0
    inline_eq_index = 0
    
    # First, replace all display equations with placeholders to handle multi-line equations
    display_pattern = r'\$\$(.*?)\$\$'
    display_placeholders = {}
    processed_content = markdown_content
    for i, match in enumerate(re.finditer(display_pattern, markdown_content, re.DOTALL)):
        placeholder = f"__DISPLAY_EQUATION_{i}__"
        display_placeholders[placeholder] = i
        processed_content = processed_content.replace(match.group(0), placeholder, 1)
    
    # Split markdown into lines and process
    lines = processed_content.split('\n')
    
    current_paragraph = None
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            if current_paragraph is None or not current_paragraph.text.strip():
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph = None
            continue
        
        # Process markdown headings
        if line.startswith('# '):
            current_paragraph = None
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            current_paragraph = None
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            current_paragraph = None
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            current_paragraph = None
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s+', line):
            # Handle numbered lists
            current_paragraph = None
            # Extract the number and text
            match = re.match(r'^(\d+)\.\s+(.*)', line)
            if match:
                text = match.group(2).strip()
                doc.add_paragraph(text, style='List Number')
        else:
            # Regular text - replace equations with OMML
            # First check for display equation placeholders
            placeholder_match = re.search(r'__DISPLAY_EQUATION_(\d+)__', line)
            if placeholder_match:
                # This line contains a display equation placeholder
                current_paragraph = None
                eq_index = int(placeholder_match.group(1))
                if eq_index < len(display_eqs):
                    eq_para = doc.add_paragraph()
                    omml = display_eqs[eq_index]
                    try:
                        if 'xmlns:m' not in omml:
                            omml_with_ns = omml.replace(
                                '<m:oMathPara>',
                                '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                            )
                        else:
                            omml_with_ns = omml
                        omml_element = parse_xml(omml_with_ns)
                        eq_para._element.append(omml_element)
                        if verbose:
                            print(f"✓ Added display equation {eq_index + 1}")
                    except Exception as e:
                        eq_para.add_run(f"[Equation - parsing error]")
                        if verbose:
                            print(f"✗ Error adding display equation {eq_index + 1}: {e}")
            else:
                # Handle inline equations
                inline_pattern = r'\$([^\$]+?)\$'
                parts = re.split(inline_pattern, line)
                
                if len(parts) > 1 and any(parts):
                    current_paragraph = None
                    p = doc.add_paragraph()
                    for i, part in enumerate(parts):
                        if i % 2 == 0:  # Text part
                            if part:
                                p.add_run(part)
                        else:  # Equation part (inline)
                            if inline_eq_index < len(inline_eqs):
                                omml = inline_eqs[inline_eq_index]
                                try:
                                    # For inline equations, extract just the <m:oMath> part
                                    # and embed it in the run
                                    if '<m:oMathPara>' in omml:
                                        # Extract the inner <m:oMath> from oMathPara
                                        start = omml.find('<m:oMath>')
                                        end = omml.find('</m:oMath>') + len('</m:oMath>')
                                        if start != -1 and end > start:
                                            omml_inner = omml[start:end]
                                        else:
                                            omml_inner = omml
                                    else:
                                        omml_inner = omml
                                    
                                    # Add namespace if needed
                                    if 'xmlns:m' not in omml_inner:
                                        omml_inner = omml_inner.replace(
                                            '<m:oMath>',
                                            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                                        )
                                    
                                    # Create a run and insert OMML
                                    run = p.add_run()
                                    omml_element = parse_xml(omml_inner)
                                    run._element.append(omml_element)
                                except Exception as e:
                                    p.add_run(f"[eq: {part}]")
                                inline_eq_index += 1
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph(line)
                    else:
                        current_paragraph.add_run('\n' + line)
    
    # Save the document
    doc.save(output_path)
    if verbose:
        print(f"Document saved to: {output_path}")
    return output_path


# ============================================================================
# MAIN PIPELINE
# ============================================================================

def markdown_to_word(markdown_file, output_file=None, verbose=True):
    """
    Convert a markdown file to a Word document with LaTeX equations rendered as OMML.
    
    This is the main entry point for the conversion pipeline.
    
    Args:
        markdown_file (str): Path to the markdown file to convert
        output_file (str): Path for the output Word document (default: markdown_file with .docx extension)
        verbose (bool): If True, print progress information
    
    Returns:
        str: Path to the created Word document
    """
    # Determine output path if not provided
    if output_file is None:
        base_path = Path(markdown_file).stem
        output_file = f"{base_path}.docx"
    
    if verbose:
        print(f"Reading markdown file: {markdown_file}")
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    if verbose:
        print(f"Extracting LaTeX equations...")
    
    # Extract equations
    equations_from_md = extract_latex_equations(markdown_content)
    
    if verbose:
        print(f"\nFound {len(equations_from_md)} LaTeX equations\n")
        for i, (eq, is_display) in enumerate(equations_from_md, 1):
            mode = "DISPLAY" if is_display else "INLINE"
            print(f"{i}. [{mode}] {eq[:60]}{'...' if len(eq) > 60 else ''}")
    
    # Convert all extracted equations to OMML
    md_omml_equations = convert_equations_to_omml(equations_from_md, verbose=verbose)
    
    # Create the comprehensive Word document
    if md_omml_equations:
        output_path = create_word_doc_from_markdown(
            markdown_content, 
            md_omml_equations,
            output_path=output_file,
            verbose=verbose
        )
        if verbose:
            print(f"\n✓ Successfully created Word document with {len(md_omml_equations)} equations")
        return output_path
    else:
        if verbose:
            print("No OMML equations to add to document")
        return None


# ============================================================================
# SCRIPT ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    import sys
    
    # Get markdown file from command-line arguments or user input
    verbose_tag = False
    if len(sys.argv) > 2:
        if sys.argv[1] == '-v':
            verbose_tag = True
        markdown_file = sys.argv[-1]
    elif len(sys.argv) > 1:
        if sys.argv[1] == '-v':
            verbose_tag = True
            markdown_file = input("Enter the path to the markdown file: ").strip()
        else:
            markdown_file = sys.argv[-1]
    else:
        markdown_file = input("Enter the path to the markdown file: ").strip()
    
    # Resolve the path (handle both relative and absolute paths)
    markdown_path = Path(markdown_file)
    
    # If relative path, check current directory first, then parent directories
    if not markdown_path.is_absolute():
        # Try current directory
        if not markdown_path.exists():
            # Try common parent directories
            alt_paths = [
                Path.cwd() / markdown_file,
                Path.cwd().parent / markdown_file,
                Path.home() / markdown_file,
            ]
            for alt_path in alt_paths:
                if alt_path.exists():
                    markdown_path = alt_path
                    break
    
    # Check if file exists
    if not markdown_path.exists():
        print(f"❌ Error: File '{markdown_file}' not found")
        print(f"   Checked: {Path.cwd() / markdown_file}")
        sys.exit(1)
    
    # Convert to absolute path for processing
    markdown_file = str(markdown_path.resolve())
    
    # Automatically generate output filename from markdown filename
    output_file = None  # This will use the default naming based on markdown_file
    
    try:
        result = markdown_to_word(markdown_file, output_file, verbose=verbose_tag)
        if result:
            print(f"\n✅ Conversion completed successfully!")
            print(f"Output file: {result}")
        else:
            print("\n❌ Conversion failed - no equations were converted")
    except Exception as e:
        print(f"\n❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
