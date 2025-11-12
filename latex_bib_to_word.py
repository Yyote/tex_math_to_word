"""
BibTeX to Word Document Converter

This module provides functionality to convert BibTeX .bib files to Word documents.
Each reference is formatted with its citation key as a label followed by the formatted
reference content (including author, title, year, journal/booktitle, etc.).
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import bibtexparser
from bibtexparser.bparser import BibTexParser
from bibtexparser.customization import convert_to_unicode


# ============================================================================
# REFERENCE FORMATTING FUNCTIONS
# ============================================================================

def format_author_list(authors):
    """
    Format the author list from BibTeX format.
    
    Args:
        authors (str): Author string from BibTeX (e.g., "A. B. Smith and J. Doe")
    
    Returns:
        str: Formatted author list
    """
    if not authors:
        return ""
    
    # Split by 'and' to get individual authors
    author_list = [a.strip() for a in authors.split(' and ')]
    
    if len(author_list) == 1:
        return author_list[0]
    elif len(author_list) == 2:
        return f"{author_list[0]} and {author_list[1]}"
    else:
        # For multiple authors: "Author1, Author2, and Author3"
        return ", ".join(author_list[:-1]) + f", and {author_list[-1]}"


def format_article(entry):
    """
    Format a journal article reference.
    
    Args:
        entry (dict): BibTeX entry dictionary
    
    Returns:
        str: Formatted reference string
    """
    parts = []
    
    # Authors
    if 'author' in entry:
        parts.append(format_author_list(entry['author']))
    
    # Title
    if 'title' in entry:
        # Remove curly braces used for capitalization protection in BibTeX
        title = entry['title'].replace('{', '').replace('}', '')
        parts.append(f'"{title}"')
    
    # Journal
    if 'journal' in entry:
        parts.append(entry['journal'])
    
    # Volume and Number
    vol_num = []
    if 'volume' in entry:
        vol_num.append(f"vol. {entry['volume']}")
    if 'number' in entry:
        vol_num.append(f"no. {entry['number']}")
    if vol_num:
        parts.append(", ".join(vol_num))
    
    # Pages
    if 'pages' in entry:
        parts.append(f"pp. {entry['pages']}")
    
    # Month and Year
    year_part = []
    if 'month' in entry:
        year_part.append(entry['month'])
    if 'year' in entry:
        year_part.append(entry['year'])
    if year_part:
        parts.append(" ".join(year_part))
    
    # DOI
    if 'doi' in entry:
        parts.append(f"doi: {entry['doi']}")
    
    return ", ".join(parts) + "."


def format_inproceedings(entry):
    """
    Format a conference proceedings reference.
    
    Args:
        entry (dict): BibTeX entry dictionary
    
    Returns:
        str: Formatted reference string
    """
    parts = []
    
    # Authors
    if 'author' in entry:
        parts.append(format_author_list(entry['author']))
    
    # Title
    if 'title' in entry:
        title = entry['title'].replace('{', '').replace('}', '')
        parts.append(f'"{title}"')
    
    # Booktitle (conference name)
    if 'booktitle' in entry:
        parts.append(f"in {entry['booktitle']}")
    
    # Year
    if 'year' in entry:
        parts.append(entry['year'])
    
    # Pages
    if 'pages' in entry:
        parts.append(f"p. {entry['pages']}")
    
    # DOI
    if 'doi' in entry:
        parts.append(f"doi: {entry['doi']}")
    
    return ", ".join(parts) + "."


def format_book(entry):
    """
    Format a book reference.
    
    Args:
        entry (dict): BibTeX entry dictionary
    
    Returns:
        str: Formatted reference string
    """
    parts = []
    
    # Authors or Editors
    if 'author' in entry:
        parts.append(format_author_list(entry['author']))
    elif 'editor' in entry:
        parts.append(format_author_list(entry['editor']) + " (Ed.)")
    
    # Title
    if 'title' in entry:
        title = entry['title'].replace('{', '').replace('}', '')
        parts.append(f"{title}")
    
    # Publisher
    if 'publisher' in entry:
        parts.append(entry['publisher'])
    
    # Year
    if 'year' in entry:
        parts.append(entry['year'])
    
    # Edition
    if 'edition' in entry:
        parts.append(f"{entry['edition']} ed.")
    
    return ", ".join(parts) + "."


def format_reference(entry):
    """
    Format a BibTeX entry based on its type.
    
    Args:
        entry (dict): BibTeX entry dictionary
    
    Returns:
        str: Formatted reference string
    """
    entry_type = entry.get('ENTRYTYPE', '').lower()
    
    if entry_type == 'article':
        return format_article(entry)
    elif entry_type in ['inproceedings', 'conference']:
        return format_inproceedings(entry)
    elif entry_type == 'book':
        return format_book(entry)
    else:
        # Generic format for other types
        parts = []
        if 'author' in entry:
            parts.append(format_author_list(entry['author']))
        if 'title' in entry:
            title = entry['title'].replace('{', '').replace('}', '')
            parts.append(f'"{title}"')
        if 'year' in entry:
            parts.append(entry['year'])
        return ", ".join(parts) + "."


# ============================================================================
# MAIN CONVERSION FUNCTIONS
# ============================================================================

def parse_bib_file(bib_filepath):
    """
    Parse a BibTeX file and return the entries.
    
    Args:
        bib_filepath (str): Path to the .bib file
    
    Returns:
        list: List of BibTeX entry dictionaries
    """
    parser = BibTexParser(common_strings=True)
    parser.customization = convert_to_unicode
    
    with open(bib_filepath, 'r', encoding='utf-8') as bib_file:
        bib_database = bibtexparser.load(bib_file, parser=parser)
    
    return bib_database.entries


def create_word_document(entries, output_filepath):
    """
    Create a Word document from BibTeX entries.
    
    Args:
        entries (list): List of BibTeX entry dictionaries
        output_filepath (str): Path to save the Word document
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Bibliography', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add each reference
    for entry in entries:
        # Get the citation key (reference label)
        cite_key = entry.get('ID', 'UNKNOWN')
        
        # Format the reference
        formatted_ref = format_reference(entry)
        
        # Add to document as a paragraph
        para = doc.add_paragraph()
        
        # Add citation key in bold
        run_label = para.add_run(f"[{cite_key}] ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        
        # Add formatted reference
        run_ref = para.add_run(formatted_ref)
        run_ref.font.size = Pt(11)
        
        # Add spacing after paragraph
        para.paragraph_format.space_after = Pt(6)
    
    # Save the document
    doc.save(output_filepath)
    print(f"Word document saved to: {output_filepath}")


def convert_bib_to_word(bib_filepath):
    """
    Convert a BibTeX file to a Word document.
    
    The Word document will be saved with the same filename as the .bib file
    in the current working directory.
    
    Args:
        bib_filepath (str): Path to the .bib file (absolute or relative)
    """
    # Convert to Path object and resolve to absolute path
    bib_path = Path(bib_filepath).resolve()
    
    # Check if file exists
    if not bib_path.exists():
        print(f"Error: File '{bib_filepath}' not found.")
        return
    
    # Check if it's a .bib file
    if bib_path.suffix.lower() != '.bib':
        print(f"Error: '{bib_filepath}' is not a .bib file.")
        return
    
    print(f"Reading BibTeX file: {bib_path}")
    
    # Parse the BibTeX file
    try:
        entries = parse_bib_file(bib_path)
        print(f"Found {len(entries)} references.")
    except Exception as e:
        print(f"Error parsing BibTeX file: {e}")
        return
    
    # Create output filename in current directory
    output_filename = bib_path.stem + '.docx'
    output_path = Path.cwd() / output_filename
    
    print(f"Creating Word document: {output_path}")
    
    # Create the Word document
    try:
        create_word_document(entries, output_path)
        print("Conversion completed successfully!")
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """
    Main function to run the BibTeX to Word converter.
    """
    # Get BibTeX file from command-line arguments or user input
    if len(sys.argv) > 1:
        bib_filepath = sys.argv[1]
    else:
        bib_filepath = input("Enter the path to the BibTeX (.bib) file: ").strip()
    
    # Check if user provided a file path
    if not bib_filepath:
        print("Error: No file path provided.")
        print("\nUsage: python latex_bib_to_word.py <path_to_bib_file>")
        print("\nExample:")
        print("  python latex_bib_to_word.py /path/to/references.bib")
        return
    
    convert_bib_to_word(bib_filepath)


if __name__ == "__main__":
    main()
