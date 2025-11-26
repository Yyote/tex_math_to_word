"""
LaTeX to Word Document Converter with LaTeX to OMML Equation Support

This module provides functionality to convert LaTeX .tex files to Word documents,
automatically converting embedded LaTeX equations to OMML (Office Math Markup Language)
format for proper rendering in Microsoft Word. It handles LaTeX environments, 
comments, and both display and inline equations.
"""

import subprocess
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from pathlib import Path


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def clean_latex_delimiters(latex_formula):
    r"""
    Clean up LaTeX delimiter commands for better OMML conversion.
    
    Handles:
    - \Bigl., \Bigr., \big., etc. (null delimiters) -> remove
    - Converts \Bigl(...) to just (...)
    - Converts \Bigr(...) to just (...)
    - \resizebox{...}{...}{$content$} -> content (strips both wrapper and $ signs)
    
    Args:
        latex_formula (str): LaTeX formula string
    
    Returns:
        str: Cleaned LaTeX formula
    """
    # Remove \resizebox{...}{...}{$...$} and keep only content without $ signs
    # This handles: \resizebox{.97\hsize}{!}{$equation$} -> equation
    def strip_resizebox_and_dollars(match):
        content = match.group(1)
        # Strip leading/trailing $ if present
        content = content.strip()
        if content.startswith('$') and content.endswith('$'):
            content = content[1:-1]
        return content
    
    formula = re.sub(r'\\resizebox\{[^}]*\}\{[^}]*\}\{(.*?)\}', strip_resizebox_and_dollars, latex_formula, flags=re.DOTALL)
    
    # Remove null delimiters: \Bigl., \Bigr., \big., \Big., \bigl., \bigr., etc.
    formula = re.sub(r'\\[Bb]ig[lr]?\.', '', formula)
    
    # Remove \Bigl and \Bigr commands (but keep the delimiter)
    # e.g., \Bigl( becomes (, \Bigr) becomes )
    formula = re.sub(r'\\[Bb]igl\s*', '', formula)
    formula = re.sub(r'\\[Bb]igr\s*', '', formula)
    formula = re.sub(r'\\big\s*', '', formula)
    formula = re.sub(r'\\Big\s*', '', formula)
    formula = re.sub(r'\\bigl\s*', '', formula)
    formula = re.sub(r'\\bigr\s*', '', formula)
    
    return formula


def latex_to_omml(latex_formula):
    """
    Convert LaTeX formula to OMML (Office Math Markup Language) using texmath.
    
    Args:
        latex_formula (str): LaTeX string (without $ or $$ delimiters)
    
    Returns:
        str or None: OMML string or None if conversion fails
    """
    try:
        # Clean up delimiter commands first
        latex_formula = clean_latex_delimiters(latex_formula)
        
        # Use full path to texmath executable
        texmath_path = os.path.expanduser("~/.local/bin/texmath")
        
        # Call texmath to convert LaTeX to OMML
        result = subprocess.run(
            [texmath_path, '--from', 'tex', '--to', 'omml'],
            input=latex_formula.encode('utf-8'),
            capture_output=True,
            timeout=5
        )
        
        if result.returncode == 0:
            return result.stdout.decode('utf-8').strip()
        else:
            print(f"Error converting formula: {result.stderr.decode('utf-8')}")
            return None
    except Exception as e:
        print(f"Exception during conversion: {e}")
        return None


def remove_latex_comments(content):
    """
    Remove LaTeX comments from content.
    
    Removes everything from % to the end of line, but preserves % inside commands.
    
    Args:
        content (str): LaTeX content as string
    
    Returns:
        str: Content with comments removed
    """
    lines = content.split('\n')
    result = []
    for line in lines:
        # Remove comments (% to end of line, but be careful about \%
        # Split by % and process
        parts = []
        in_verb = False
        i = 0
        while i < len(line):
            if i > 0 and line[i] == '%' and line[i-1] != '\\':
                # Found a comment
                break
            parts.append(line[i])
            i += 1
        result.append(''.join(parts).rstrip())
    
    # Remove empty lines at the end of content
    while result and not result[-1]:
        result.pop()
    
    return '\n'.join(result)


def skip_latex_preamble(content):
    r"""
    Skip LaTeX document preamble (everything before \begin{document} or \section).
    
    Args:
        content (str): LaTeX content as string
    
    Returns:
        str: Content starting from main content
    """
    # Look for \begin{document}
    doc_start = content.find(r'\begin{document}')
    if doc_start != -1:
        return content[doc_start + len(r'\begin{document}'):]
    
    # If no \begin{document}, look for first \section, \chapter, etc.
    section_patterns = [r'\section{', r'\chapter{', r'\part{', r'\subsection{']
    earliest_pos = len(content)
    
    for pattern in section_patterns:
        pos = content.find(pattern)
        if pos != -1 and pos < earliest_pos:
            earliest_pos = pos
    
    if earliest_pos < len(content):
        return content[earliest_pos:]
    
    return content


def extract_latex_equations(content):
    r"""
    Extract LaTeX equations from LaTeX content.
    
    Extracts both display equations (\begin{equation}...\end{equation}, $$...$$)
    and inline equations ($...$).
    
    Args:
        content (str): LaTeX content as string
    
    Returns:
        list: List of tuples (equation, is_display_mode, label) where is_display_mode is bool and label is str or None
    """
    equations = []
    
    # Helper function to extract label(s) from equation
    def extract_label(eq_text):
        # Find all labels (for multi-line equations like align with multiple labels)
        label_matches = re.findall(r'\\label\{([^}]*)\}', eq_text)
        if label_matches:
            # Join multiple labels with comma
            label = ', '.join(label_matches)
            # Remove all labels from equation text
            eq_text = re.sub(r'\\label\{[^}]*\}', '', eq_text).strip()
            return eq_text, label
        return eq_text, None
    
    # First extract equation environments (display)
    # Patterns: \begin{equation*?}...\end{equation*?}, \begin{align*?}...\end{align*?}
    # For align/gather/etc, we need to keep the environment wrapper for texmath
    env_patterns_no_wrapper = [
        (r'\\begin\{equation\*?\}(.*?)\\end\{equation\*?\}', 'equation'),
    ]
    
    env_patterns_with_wrapper = [
        (r'\\begin\{(align\*?)\}(.*?)\\end\{\1\}', 'align'),
        (r'\\begin\{(gather\*?)\}(.*?)\\end\{\1\}', 'gather'),
        (r'\\begin\{(multline\*?)\}(.*?)\\end\{\1\}', 'multline'),
        (r'\\begin\{(split)\}(.*?)\\end\{\1\}', 'split'),
    ]
    
    # Extract equations that don't need wrapper (like equation environment)
    for pattern, env_name in env_patterns_no_wrapper:
        for match in re.finditer(pattern, content, re.DOTALL):
            eq = match.group(1).strip()
            if eq:
                eq_clean, label = extract_label(eq)
                equations.append((eq_clean, True, label))
    
    # Extract equations that need wrapper (like align, gather, etc.)
    for pattern, env_name in env_patterns_with_wrapper:
        for match in re.finditer(pattern, content, re.DOTALL):
            env_type = match.group(1)
            eq_content = match.group(2).strip()
            if eq_content:
                # Extract label before adding wrapper back
                eq_clean, label = extract_label(eq_content)
                # Keep the environment wrapper for texmath to parse alignment correctly
                eq_with_wrapper = f'\\begin{{{env_type}}}\n{eq_clean}\n\\end{{{env_type}}}'
                equations.append((eq_with_wrapper, True, label))
    
    # Remove equation environments from content for next extraction
    temp_content = content
    for pattern, _ in env_patterns_no_wrapper:
        temp_content = re.sub(pattern, '', temp_content, flags=re.DOTALL)
    for pattern, _ in env_patterns_with_wrapper:
        temp_content = re.sub(pattern, '', temp_content, flags=re.DOTALL)
    
    # Extract display equations ($$...$$)
    display_pattern = r'\$\$(.*?)\$\$'
    for match in re.finditer(display_pattern, temp_content, re.DOTALL):
        eq = match.group(1).strip()
        if eq:
            eq_clean, label = extract_label(eq)
            equations.append((eq_clean, True, label))  # True = display mode
    
    # Remove display equations from content for inline extraction
    temp_content = re.sub(display_pattern, '', temp_content, flags=re.DOTALL)
    
    # Extract inline equations ($...$)
    inline_pattern = r'\$([^\$]+?)\$'
    for match in re.finditer(inline_pattern, temp_content):
        eq = match.group(1).strip()
        if eq and '\n' not in eq:  # Skip if multiline
            eq_clean, label = extract_label(eq)
            equations.append((eq_clean, False, label))  # False = inline mode
    
    return equations


def convert_equations_to_omml(equations_list, verbose=False):
    """
    Convert a list of LaTeX equations to OMML format.
    
    Args:
        equations_list (list): List of tuples (equation, is_display, label)
        verbose (bool): If True, print progress information
    
    Returns:
        list: List of tuples (omml, is_display, label) for successfully converted equations
    """
    omml_equations = []
    
    if verbose:
        print("Converting extracted equations to OMML:\n")
    
    for i, (latex_eq, is_display, label) in enumerate(equations_list, 1):
        mode = "DISPLAY" if is_display else "INLINE"
        label_info = f" [Label: {label}]" if label else ""
        if verbose:
            print(f"Converting equation {i} [{mode}]{label_info}: {latex_eq[:40]}...")
        
        omml = latex_to_omml(latex_eq)
        if omml:
            omml_equations.append((omml, is_display, label))
            if verbose:
                print(f"  ✓ Success")
        else:
            if verbose:
                print(f"  ✗ Failed")
    
    if verbose:
        print(f"\nSuccessfully converted {len(omml_equations)} out of {len(equations_list)} equations")
    
    return omml_equations


# ============================================================================
# LaTeX PROCESSING
# ============================================================================

def process_latex_structure(content):
    r"""
    Convert basic LaTeX structure to plain text with markers for processing.
    
    Handles:
    - Section headers (\section, \subsection, etc.)
    - Basic text processing
    - Preserves equations (marked with placeholders)
    - Extracts figure captions and labels while omitting figure contents
    - Preserves inline equations in figure captions for later conversion
    
    Args:
        content (str): LaTeX content
    
    Returns:
        str: Processed content with LaTeX commands removed/converted
    """
    # Handle \texorpdfstring{arg1}{arg2} - keep only arg2 (do this EARLY before other processing)
    # Need to handle nested braces properly
    def replace_texorpdfstring(content):
        while r'\texorpdfstring{' in content:
            match = re.search(r'\\texorpdfstring\{', content)
            if not match:
                break
            
            start_pos = match.end()
            # Find the end of first argument
            brace_count = 1
            pos = start_pos
            while pos < len(content) and brace_count > 0:
                if content[pos] == '{' and (pos == 0 or content[pos-1] != '\\'):
                    brace_count += 1
                elif content[pos] == '}' and (pos == 0 or content[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            
            if brace_count != 0:
                break  # Malformed
            
            # Now pos is at the closing brace of arg1, next should be {arg2}
            if pos >= len(content) or content[pos] != '{':
                break
            
            # Find arg2
            start_arg2 = pos + 1
            brace_count = 1
            pos = start_arg2
            while pos < len(content) and brace_count > 0:
                if content[pos] == '{' and (pos == 0 or content[pos-1] != '\\'):
                    brace_count += 1
                elif content[pos] == '}' and (pos == 0 or content[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            
            if brace_count != 0:
                break
            
            # Extract arg2 and replace the whole \texorpdfstring{arg1}{arg2}
            arg2 = content[start_arg2:pos-1]
            content = content[:match.start()] + arg2 + content[pos:]
        
        return content
    
    content = replace_texorpdfstring(content)
    
    # Strip \resizebox{...}{...}{content} and keep only content (without $ if present)
    # Use brace counting to handle nested braces properly
    # Note: This function is now redundant since we strip resizebox earlier in latex_to_word()
    # but keeping it here for any edge cases in process_latex_structure
    def strip_all_resizebox(text):
        iteration = 0
        while r'\resizebox{' in text:
            iteration += 1
            if iteration > 100:  # Safety check to prevent infinite loops
                break
                
            match = re.search(r'\\resizebox\{', text)
            if not match:
                break
            
            # The match.end() position is right after the opening brace of first arg
            # Pattern: \resizebox{width}{height}{content}
            pos = match.end() - 1  # Back to the '{' of first argument
            
            # Skip first two arguments {width}{height}
            for arg_num in range(2):
                if pos >= len(text) or text[pos] != '{':
                    break
                brace_count = 1
                pos += 1  # Move past the opening '{'
                while pos < len(text) and brace_count > 0:
                    if text[pos] == '{' and (pos == 0 or text[pos-1] != '\\'):
                        brace_count += 1
                    elif text[pos] == '}' and (pos == 0 or text[pos-1] != '\\'):
                        brace_count -= 1
                    pos += 1
            
            # Now pos should be at the opening '{' of the third argument (the content)
            if pos >= len(text) or text[pos] != '{':
                break
            
            start_content = pos + 1
            brace_count = 1
            pos = start_content
            while pos < len(text) and brace_count > 0:
                if text[pos] == '{' and (pos == 0 or text[pos-1] != '\\'):
                    brace_count += 1
                elif text[pos] == '}' and (pos == 0 or text[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            
            if brace_count != 0:
                break
            
            # Extract content and remove surrounding $ if present
            content_inner = text[start_content:pos-1].strip()
            if content_inner.startswith('$') and content_inner.endswith('$'):
                content_inner = content_inner[1:-1]
            
            # Replace the whole \resizebox{...}{...}{...} with just the content
            text = text[:match.start()] + content_inner + text[pos:]
        
        return text
    
    content = strip_all_resizebox(content)
    
    # Extract figure captions and labels from figure environments
    def extract_figure_info(match):
        fig_content = match.group(1)
        
        # Extract caption - need to handle nested braces properly
        caption_match = re.search(r'\\caption\{', fig_content)
        caption_text = ""
        if caption_match:
            # Find the matching closing brace for \caption{
            start_pos = caption_match.end()
            brace_count = 1
            pos = start_pos
            while pos < len(fig_content) and brace_count > 0:
                if fig_content[pos] == '{' and (pos == 0 or fig_content[pos-1] != '\\'):
                    brace_count += 1
                elif fig_content[pos] == '}' and (pos == 0 or fig_content[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            if brace_count == 0:
                caption_text = fig_content[start_pos:pos-1]
        
        label = re.search(r'\\label\{([^}]*)\}', fig_content)
        
        output = []
        if label:
            output.append(f"[Figure: {label.group(1)}]")
        if caption_text:
            # Encapsulate caption in brackets and keep inline equations intact (they'll be processed later)
            output.append(f"[Caption: {caption_text}]")
        
        if output:
            return '\n' + '\n'.join(output) + '\n'
        else:
            return '[Figure omitted - not supported in conversion]\n'
    
    # Handle figure environments - extract captions and labels
    content = re.sub(
        r'\\begin\{figure\*?\}(.*?)\\end\{figure\*?\}',
        extract_figure_info,
        content,
        flags=re.DOTALL
    )
    
    # Handle table environments - extract table content and captions
    def extract_table_info(match):
        table_content = match.group(1)
        
        # Extract caption
        caption_match = re.search(r'\\caption\{', table_content)
        caption_text = ""
        if caption_match:
            start_pos = caption_match.end()
            brace_count = 1
            pos = start_pos
            while pos < len(table_content) and brace_count > 0:
                if table_content[pos] == '{' and (pos == 0 or table_content[pos-1] != '\\'):
                    brace_count += 1
                elif table_content[pos] == '}' and (pos == 0 or table_content[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            if brace_count == 0:
                caption_text = table_content[start_pos:pos-1]
        
        # Extract label
        label = re.search(r'\\label\{([^}]*)\}', table_content)
        
        # Extract tabular/tabularx content with proper brace counting
        tabular_start = re.search(r'\\begin\{(tabular[x]?)\}', table_content)
        table_data = None
        
        if tabular_start:
            env_name = tabular_start.group(1)
            pos = tabular_start.end()
            
            # For tabularx, skip two arguments: {width}{column_spec}
            # For tabular, skip one argument: {column_spec}
            num_args_to_skip = 2 if env_name == 'tabularx' else 1
            
            for _ in range(num_args_to_skip):
                if pos < len(table_content) and table_content[pos] == '{':
                    brace_count = 1
                    pos += 1
                    while pos < len(table_content) and brace_count > 0:
                        if table_content[pos] == '{' and (pos == 0 or table_content[pos-1] != '\\'):
                            brace_count += 1
                        elif table_content[pos] == '}' and (pos == 0 or table_content[pos-1] != '\\'):
                            brace_count -= 1
                        pos += 1
            
            # Now extract content until \end{tabular}
            start_pos = pos
            end_pattern = f'\\end{{{env_name}}}'
            end_pos = table_content.find(end_pattern, start_pos)
            
            if end_pos != -1:
                table_data = table_content[start_pos:end_pos].strip()
        
        output = []
        if label:
            output.append(f"[Table: {label.group(1)}]")
        if caption_text:
            output.append(f"[Caption: {caption_text}]")
        
        if table_data:
            # Mark the table content for processing
            output.append(f"__TABLE_START__\n{table_data}\n__TABLE_END__")
        else:
            output.append('[Table content could not be extracted]')
        
        if output:
            return '\n' + '\n'.join(output) + '\n'
        else:
            return '[Table omitted - not supported in conversion]\n'
    
    # Handle table environments
    content = re.sub(
        r'\\begin\{table\*?\}(.*?)\\end\{table\*?\}',
        extract_table_info,
        content,
        flags=re.DOTALL
    )
    
    # Handle itemize environments - convert to marked list items
    def process_itemize(match):
        items_content = match.group(1)
        # Extract individual \item entries
        items = re.findall(r'\\item\s+(.*?)(?=\\item|$)', items_content, re.DOTALL)
        result = []
        for item in items:
            item = item.strip()
            if item:
                result.append(f'__BULLET_ITEM__{item}')
        return '\n' + '\n'.join(result) + '\n'
    
    content = re.sub(
        r'\\begin\{itemize\}(.*?)\\end\{itemize\}',
        process_itemize,
        content,
        flags=re.DOTALL
    )
    
    # Handle enumerate environments - convert to marked numbered list items
    def process_enumerate(match):
        items_content = match.group(1)
        # Extract individual \item entries
        items = re.findall(r'\\item\s+(.*?)(?=\\item|$)', items_content, re.DOTALL)
        result = []
        for item in items:
            item = item.strip()
            if item:
                result.append(f'__NUMBERED_ITEM__{item}')
        return '\n' + '\n'.join(result) + '\n'
    
    content = re.sub(
        r'\\begin\{enumerate\}(.*?)\\end\{enumerate\}',
        process_enumerate,
        content,
        flags=re.DOTALL
    )
    
    # Remove \label{...} and convert reference commands
    content = re.sub(r'\\label\{[^}]*\}', '', content)
    # Convert \ref{label} to [label]
    def replace_ref(match):
        label = match.group(1)
        return f'[{label}]'
    content = re.sub(r'\\ref\{([^}]*)\}', replace_ref, content)
    # Keep citation labels: \cite{label1,label2} -> [label1,label2]
    def replace_cite(match):
        labels = match.group(1)
        return f'[{labels}]'
    content = re.sub(r'\\cite\{([^}]*)\}', replace_cite, content)
    
    # Remove \reffig{...} and \refeqn{...}
    content = re.sub(r'\\reffig\{[^}]*\}', 'Fig.', content)
    content = re.sub(r'\\refeqn\{[^}]*\}', 'Eq.', content)
    
    # Handle text formatting commands
    content = re.sub(r'\\textbf\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\textit\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\texttt\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\emph\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\text\{([^}]*)\}', r'\1', content)
    
    # Handle subscript and superscript in text (not math mode)
    content = re.sub(r'\\textsubscript\{([^}]*)\}', r'__SUB__\1__/SUB__', content)
    content = re.sub(r'\\textsuperscript\{([^}]*)\}', r'__SUP__\1__/SUP__', content)
    
    # Handle other common commands
    content = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\mathbb\{([^}]*)\}', r'\1', content)
    
    # Replace LaTeX non-breaking space ~ with regular space
    content = re.sub(r'~', ' ', content)
    
    # Handle line breaks (but protect table content)
    # Split by table markers and only process non-table parts
    parts = re.split(r'(__TABLE_START__.*?__TABLE_END__)', content, flags=re.DOTALL)
    for i in range(len(parts)):
        if not parts[i].startswith('__TABLE_START__'):
            parts[i] = re.sub(r'\\\\', '\n', parts[i])
    content = ''.join(parts)
    
    # Replace escaped percentage signs \% with %
    content = re.sub(r'\\%', '%', content)
    
    # Clean up extra whitespace
    lines = content.split('\n')
    lines = [line.strip() for line in lines]
    content = '\n'.join(lines)
    
    return content


# ============================================================================
# DOCUMENT GENERATION
# ============================================================================

def add_formatted_text(paragraph, text):
    """
    Add text to a paragraph with support for subscript and superscript formatting.
    
    Handles __SUB__text__/SUB__ and __SUP__text__/SUP__ markers.
    
    Args:
        paragraph: The paragraph object to add text to
        text (str): Text with formatting markers
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Pattern to match subscript and superscript markers
    pattern = r'__SUB__([^_]+?)__/SUB__|__SUP__([^_]+?)__/SUP__'
    
    parts = re.split(pattern, text)
    
    for i, part in enumerate(parts):
        if part is None:
            continue
        if i % 3 == 0:  # Regular text
            if part:
                paragraph.add_run(part)
        elif i % 3 == 1:  # Subscript text (from __SUB__)
            if part:
                run = paragraph.add_run(part)
                run.font.subscript = True
        elif i % 3 == 2:  # Superscript text (from __SUP__)
            if part:
                run = paragraph.add_run(part)
                run.font.superscript = True


def parse_latex_table(table_text):
    """
    Parse LaTeX table content and return a list of rows with proper multicolumn handling.
    
    Args:
        table_text (str): Content between \\begin{tabular} and \\end{tabular}
    
    Returns:
        list: List of rows, where each row is a list of (content, colspan) tuples
    """
    # Remove \hline, \cline, and other formatting commands
    table_text = re.sub(r'\\hline', '', table_text)
    table_text = re.sub(r'\\cline\{[^}]*\}', '', table_text)
    table_text = re.sub(r'\\centering', '', table_text)
    
    # Split by \\ to get rows (but not \\\\)
    rows = re.split(r'\\\\(?!\\)', table_text)
    
    table_data = []
    for row in rows:
        row = row.strip()
        if not row:
            continue
        
        # Helper function to extract content from braces
        def extract_braced_content(text, start_pos):
            """Extract content within braces starting at start_pos (which should be at '{')"""
            if start_pos >= len(text) or text[start_pos] != '{':
                return "", start_pos
            
            content = ""
            pos = start_pos + 1
            brace_count = 1
            
            while pos < len(text) and brace_count > 0:
                if text[pos] == '\\' and pos + 1 < len(text):
                    content += text[pos:pos+2]
                    pos += 2
                elif text[pos] == '{':
                    brace_count += 1
                    content += text[pos]
                    pos += 1
                elif text[pos] == '}':
                    brace_count -= 1
                    if brace_count > 0:
                        content += text[pos]
                    pos += 1
                else:
                    content += text[pos]
                    pos += 1
            
            return content.strip(), pos
        
        # Parse cells with multicolumn support
        cells = []
        i = 0
        current_cell = ""
        brace_count = 0
        
        while i < len(row):
            # Check for \multicolumn{n}{format}{content}
            if row[i:].startswith('\\multicolumn'):
                i += 12
                # Skip whitespace
                while i < len(row) and row[i] in ' \t':
                    i += 1
                
                # Extract colspan (first argument)
                colspan_str, i = extract_braced_content(row, i)
                colspan = int(colspan_str) if colspan_str.isdigit() else 1
                
                # Skip format (second argument)
                _, i = extract_braced_content(row, i)
                
                # Extract content (third argument)
                content, i = extract_braced_content(row, i)
                
                cells.append((content, colspan))
                continue
            
            # Check for \multirow{n}{width}{content}
            elif row[i:].startswith('\\multirow'):
                i += 9
                # Skip whitespace
                while i < len(row) and row[i] in ' \t':
                    i += 1
                
                # Skip first two arguments
                _, i = extract_braced_content(row, i)
                _, i = extract_braced_content(row, i)
                
                # Extract content (third argument)
                content, i = extract_braced_content(row, i)
                
                cells.append((content, 1))
                continue
            
            # Check for & separator
            elif row[i] == '&' and brace_count == 0:
                if current_cell.strip():
                    cells.append((current_cell.strip(), 1))
                current_cell = ""
                i += 1
                continue
            
            # Track braces for non-command content
            elif row[i] == '\\' and i + 1 < len(row):
                current_cell += row[i:i+2]
                i += 2
            elif row[i] == '{':
                brace_count += 1
                current_cell += row[i]
                i += 1
            elif row[i] == '}':
                brace_count -= 1
                current_cell += row[i]
                i += 1
            else:
                current_cell += row[i]
                i += 1
        
        # Add the last cell
        if current_cell.strip():
            cells.append((current_cell.strip(), 1))
        
        if cells:
            table_data.append(cells)
    
    return table_data


def add_paragraph_with_equations(doc, text, inline_eqs, inline_eq_index, verbose=False):
    """
    Add a paragraph to the document, replacing inline equation placeholders with OMML.
    Also handles subscript/superscript formatting.
    
    Args:
        doc: The Document object
        text (str): Text containing $..$ inline equations and formatting markers
        inline_eqs (list): List of OMML inline equations
        inline_eq_index (int): Current index in inline_eqs list
        verbose (bool): Print debug info
    
    Returns:
        int: Updated inline_eq_index
    """
    inline_pattern = r'\$([^\$]+?)\$'
    parts = re.split(inline_pattern, text)
    
    if len(parts) > 1 and any(parts):
        p = doc.add_paragraph()
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Text part
                if part:
                    # Handle subscripts and superscripts in text
                    add_formatted_text(p, part)
            else:  # Equation part (inline)
                if inline_eq_index < len(inline_eqs):
                    omml = inline_eqs[inline_eq_index]
                    try:
                        # For inline equations, extract just the <m:oMath> part
                        if '<m:oMathPara>' in omml:
                            # Extract the inner <m:oMath> from oMathPara
                            start = omml.find('<m:oMath>')
                            end = omml.find('</m:oMath>') + len('</m:oMath>')
                            if start != -1 and end > start:
                                omml_inner = omml[start:end]
                            else:
                                omml_inner = omml
                        else:
                            omml_inner = omml
                        
                        # Add namespace if needed
                        if 'xmlns:m' not in omml_inner:
                            omml_inner = omml_inner.replace(
                                '<m:oMath>',
                                '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                            )
                        
                        # Create a run and insert OMML
                        run = p.add_run()
                        omml_element = parse_xml(omml_inner)
                        run._element.append(omml_element)
                        if verbose:
                            print(f"  ✓ Added inline equation {inline_eq_index + 1}")
                    except Exception as e:
                        p.add_run(f"[eq: {part}]")
                        if verbose:
                            print(f"  ✗ Failed to add inline equation: {e}")
                    inline_eq_index += 1
                else:
                    # No more OMML equations available, add as text placeholder
                    p.add_run(f"[eq: {part}]")
                    if verbose:
                        print(f"  ⚠ Warning: Not enough OMML equations for inline equation")
    else:
        p = doc.add_paragraph()
        add_formatted_text(p, text)
    
    return inline_eq_index


def create_word_doc_from_latex(latex_content, omml_equations_data, output_path="output.docx", verbose=False):
    """
    Create a Word document from LaTeX content with OMML equations.
    
    Args:
        latex_content (str): Full LaTeX text
        omml_equations_data (list): List of (omml, is_display) tuples
        output_path (str): Path where the document will be saved
        verbose (bool): If True, print progress information
    
    Returns:
        str: Path to the created document
    """
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    # Process LaTeX structure to plain text
    processed_content = process_latex_structure(latex_content)
    
    # Create a new Document
    doc = Document()
    
    # Create a mapping of equation positions
    display_eqs = [(omml, label) for omml, is_display, label in omml_equations_data if is_display]
    inline_eqs = [omml for omml, is_display, label in omml_equations_data if not is_display]
    
    # Index to track which OMML equation we're on
    display_eq_index = 0
    inline_eq_index = 0
    
    # First, replace all display equations with placeholders
    env_patterns = [
        r'\\begin\{equation\*?\}(.*?)\\end\{equation\*?\}',
        r'\\begin\{align\*?\}(.*?)\\end\{align\*?\}',
        r'\\begin\{gather\*?\}(.*?)\\end\{gather\*?\}',
        r'\\begin\{multline\*?\}(.*?)\\end\{multline\*?\}',
        r'\\begin\{split\}(.*?)\\end\{split\}',
    ]
    
    display_placeholders = {}
    placeholder_counter = 0
    for pattern in env_patterns:
        for match in re.finditer(pattern, processed_content, re.DOTALL):
            placeholder = f"__DISPLAY_EQUATION_{placeholder_counter}__"
            display_placeholders[placeholder] = placeholder_counter
            processed_content = processed_content.replace(match.group(0), placeholder, 1)
            placeholder_counter += 1
    
    # Replace $$...$$ equations
    display_pattern = r'\$\$(.*?)\$\$'
    for i, match in enumerate(re.finditer(display_pattern, processed_content, re.DOTALL)):
        placeholder = f"__DISPLAY_EQUATION_{placeholder_counter}__"
        display_placeholders[placeholder] = placeholder_counter
        processed_content = processed_content.replace(match.group(0), placeholder, 1)
        placeholder_counter += 1
    
    # Split into lines and process
    lines = processed_content.split('\n')
    
    current_paragraph = None
    in_table = False
    table_content = []
    
    for line in lines:
        line = line.rstrip()
        
        # Handle table markers
        if line == '__TABLE_START__':
            in_table = True
            table_content = []
            current_paragraph = None
            continue
        elif line == '__TABLE_END__':
            in_table = False
            if table_content:
                # Parse and create table
                joined_content = '\n'.join(table_content)
                table_rows = parse_latex_table(joined_content)
                if table_rows:
                    # Calculate actual number of columns (accounting for colspan)
                    max_cols = max(sum(colspan for _, colspan in row) for row in table_rows)
                    
                    # Create Word table
                    word_table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    word_table.style = 'Table Grid'
                    
                    # Fill in the table
                    for i, row_data in enumerate(table_rows):
                        col_index = 0
                        for cell_text, colspan in row_data:
                            if col_index < max_cols:
                                cell = word_table.rows[i].cells[col_index]
                                
                                # Merge cells if colspan > 1
                                if colspan > 1 and col_index + colspan <= max_cols:
                                    merge_cell = word_table.rows[i].cells[col_index + colspan - 1]
                                    cell.merge(merge_cell)
                                
                                # Clear the cell first
                                cell.text = ''
                                paragraph = cell.paragraphs[0]
                                
                                # Process cell text with inline equations
                                # Find all inline math expressions
                                inline_math_pattern = r'\$([^\$]+)\$'
                                last_end = 0
                                
                                for match in re.finditer(inline_math_pattern, cell_text):
                                    # Add text before the equation
                                    if match.start() > last_end:
                                        text_before = cell_text[last_end:match.start()]
                                        # Clean LaTeX commands
                                        text_before = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text_before)
                                        text_before = re.sub(r'\\pm', '±', text_before)
                                        paragraph.add_run(text_before)
                                    
                                    # Convert and add the equation
                                    latex_expr = match.group(1)
                                    omml = latex_to_omml(latex_expr)
                                    if omml:
                                        try:
                                            # For inline equations, extract just the <m:oMath> part
                                            if '<m:oMathPara>' in omml:
                                                # Extract the inner <m:oMath> from oMathPara
                                                start = omml.find('<m:oMath>')
                                                end = omml.find('</m:oMath>') + len('</m:oMath>')
                                                if start != -1 and end > start:
                                                    omml_inner = omml[start:end]
                                                else:
                                                    omml_inner = omml
                                            else:
                                                omml_inner = omml
                                            
                                            # Add namespace if needed
                                            if 'xmlns:m' not in omml_inner:
                                                omml_inner = omml_inner.replace(
                                                    '<m:oMath>',
                                                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                                                )
                                            
                                            # Create a run and insert OMML
                                            run = paragraph.add_run()
                                            omml_element = parse_xml(omml_inner)
                                            run._element.append(omml_element)
                                        except Exception as e:
                                            # Fallback to plain text
                                            paragraph.add_run(latex_expr)
                                            if verbose:
                                                print(f"  ✗ Failed to add table cell equation: {e}")
                                    else:
                                        # Fallback to plain text
                                        paragraph.add_run(latex_expr)
                                    
                                    last_end = match.end()
                                
                                # Add remaining text after last equation
                                if last_end < len(cell_text):
                                    text_after = cell_text[last_end:]
                                    # Clean LaTeX commands
                                    text_after = re.sub(r'\\textbf\{([^}]*)\}', r'\1', text_after)
                                    text_after = re.sub(r'\\pm', '±', text_after)
                                    paragraph.add_run(text_after)
                                
                                col_index += colspan
                    
                    if verbose:
                        print(f"✓ Added table with {len(table_rows)} rows and {max_cols} columns")
            table_content = []
            continue
        elif in_table:
            table_content.append(line)
            continue
        
        if not line:
            if current_paragraph is None or not current_paragraph.text.strip():
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph = None
            continue
        
        # Process LaTeX section commands
        if line.startswith(r'\chapter{'):
            current_paragraph = None
            chapter_text = re.search(r'\\chapter\{([^}]*)\}', line)
            if chapter_text:
                heading = chapter_text.group(1)
                # Use level 0 for chapter (larger than section)
                doc.add_heading(heading, level=0)
        elif line.startswith(r'\section{'):
            current_paragraph = None
            section_text = re.search(r'\\section\{([^}]*)\}', line)
            if section_text:
                heading = section_text.group(1)
                doc.add_heading(heading, level=1)
        elif line.startswith(r'\subsection{'):
            current_paragraph = None
            subsec_text = re.search(r'\\subsection\{([^}]*)\}', line)
            if subsec_text:
                heading = subsec_text.group(1)
                doc.add_heading(heading, level=2)
        elif line.startswith(r'\subsubsection{'):
            current_paragraph = None
            subsubsec_text = re.search(r'\\subsubsection\{([^}]*)\}', line)
            if subsubsec_text:
                heading = subsubsec_text.group(1)
                doc.add_heading(heading, level=3)
        elif line.startswith('__BULLET_ITEM__'):
            # Handle bullet list items
            current_paragraph = None
            item_text = line[len('__BULLET_ITEM__'):]
            inline_eq_index = add_paragraph_with_equations(
                doc, item_text, inline_eqs, inline_eq_index, verbose=verbose
            )
            # Add bullet style to the last paragraph
            doc.paragraphs[-1].style = 'List Bullet'
        elif line.startswith('__NUMBERED_ITEM__'):
            # Handle numbered list items
            current_paragraph = None
            item_text = line[len('__NUMBERED_ITEM__'):]
            inline_eq_index = add_paragraph_with_equations(
                doc, item_text, inline_eqs, inline_eq_index, verbose=verbose
            )
            # Add numbered list style to the last paragraph
            doc.paragraphs[-1].style = 'List Number'
        else:
            # Regular text - replace equations with OMML
            # First check for display equation placeholders
            placeholder_match = re.search(r'__DISPLAY_EQUATION_(\d+)__', line)
            if placeholder_match:
                # This line contains a display equation placeholder
                current_paragraph = None
                eq_index = int(placeholder_match.group(1))
                if eq_index < len(display_eqs):
                    eq_para = doc.add_paragraph()
                    omml, label = display_eqs[eq_index]
                    try:
                        if 'xmlns:m' not in omml:
                            omml_with_ns = omml.replace(
                                '<m:oMathPara>',
                                '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                            )
                        else:
                            omml_with_ns = omml
                        omml_element = parse_xml(omml_with_ns)
                        eq_para._element.append(omml_element)
                        if verbose:
                            print(f"✓ Added display equation {eq_index + 1}")
                        
                        # Add equation label as a separate paragraph below the equation if it exists
                        if label:
                            label_para = doc.add_paragraph(f"[{label}]")
                            if verbose:
                                print(f"  Added label: [{label}]")
                    except Exception as e:
                        eq_para.add_run(f"[Equation - parsing error]")
                        if verbose:
                            print(f"✗ Error adding display equation {eq_index + 1}: {e}")
            else:
                # Handle inline equations using the helper function
                current_paragraph = None
                inline_eq_index = add_paragraph_with_equations(
                    doc, line, inline_eqs, inline_eq_index, verbose=verbose
                )
    
    # Save the document
    doc.save(output_path)
    if verbose:
        print(f"Document saved to: {output_path}")
    return output_path


# ============================================================================
# MAIN PIPELINE
# ============================================================================

def latex_to_word(latex_file, output_file=None, verbose=True):
    """
    Convert a LaTeX file to a Word document with LaTeX equations rendered as OMML.
    
    This is the main entry point for the conversion pipeline.
    
    Args:
        latex_file (str): Path to the LaTeX file to convert
        output_file (str): Path for the output Word document (default: latex_file with .docx extension)
        verbose (bool): If True, print progress information
    
    Returns:
        str: Path to the created Word document
    """
    # Determine output path if not provided
    if output_file is None:
        base_path = Path(latex_file).stem
        output_file = f"{base_path}.docx"
    
    if verbose:
        print(f"Reading LaTeX file: {latex_file}")
    
    # Read the LaTeX file
    with open(latex_file, 'r', encoding='utf-8') as f:
        latex_content = f.read()
    
    if verbose:
        print(f"Removing comments...")
    
    # Remove LaTeX comments
    latex_content = remove_latex_comments(latex_content)
    
    if verbose:
        print(f"Skipping preamble (if any)...")
    
    # Skip preamble
    latex_content = skip_latex_preamble(latex_content)
    
    # Strip \resizebox commands BEFORE extracting equations
    # This is critical because \resizebox{...}{...}{$...$} will confuse the equation extractor
    if verbose and r'\resizebox' in latex_content:
        print(f"Stripping \\resizebox commands...")
    
    def strip_resizebox_early(text):
        """Strip \resizebox{...}{...}{content} and remove $ delimiters."""
        iteration = 0
        while r'\resizebox{' in text:
            iteration += 1
            if iteration > 100:
                break
            match = re.search(r'\\resizebox\{', text)
            if not match:
                break
            
            pos = match.end() - 1
            for _ in range(2):
                if pos >= len(text) or text[pos] != '{':
                    break
                brace_count = 1
                pos += 1
                while pos < len(text) and brace_count > 0:
                    if text[pos] == '{' and (pos == 0 or text[pos-1] != '\\'):
                        brace_count += 1
                    elif text[pos] == '}' and (pos == 0 or text[pos-1] != '\\'):
                        brace_count -= 1
                    pos += 1
            
            if pos >= len(text) or text[pos] != '{':
                break
            
            start_content = pos + 1
            brace_count = 1
            pos = start_content
            while pos < len(text) and brace_count > 0:
                if text[pos] == '{' and (pos == 0 or text[pos-1] != '\\'):
                    brace_count += 1
                elif text[pos] == '}' and (pos == 0 or text[pos-1] != '\\'):
                    brace_count -= 1
                pos += 1
            
            if brace_count != 0:
                break
            
            content_inner = text[start_content:pos-1].strip()
            if content_inner.startswith('$') and content_inner.endswith('$'):
                content_inner = content_inner[1:-1]
            
            text = text[:match.start()] + content_inner + text[pos:]
        
        return text
    
    latex_content = strip_resizebox_early(latex_content)
    
    if verbose:
        print(f"Extracting LaTeX equations...")
    
    # Extract equations
    equations_from_tex = extract_latex_equations(latex_content)
    
    if verbose:
        print(f"\nFound {len(equations_from_tex)} LaTeX equations\n")
        for i, (eq, is_display, label) in enumerate(equations_from_tex, 1):
            mode = "DISPLAY" if is_display else "INLINE"
            eq_preview = eq.replace('\n', ' ')[:60]
            label_str = f" [Label: {label}]" if label else ""
            print(f"{i}. [{mode}]{label_str} {eq_preview}{'...' if len(eq) > 60 else ''}")
    
    # Convert all extracted equations to OMML
    tex_omml_equations = convert_equations_to_omml(equations_from_tex, verbose=verbose)
    
    # Create the comprehensive Word document
    if tex_omml_equations or len(equations_from_tex) == 0:
        output_path = create_word_doc_from_latex(
            latex_content, 
            tex_omml_equations,
            output_path=output_file,
            verbose=verbose
        )
        if verbose:
            if tex_omml_equations:
                print(f"\n✓ Successfully created Word document with {len(tex_omml_equations)} equations")
            else:
                print(f"\n✓ Successfully created Word document (no equations to convert)")
        return output_path
    else:
        if verbose:
            print("No OMML equations to add to document")
        return None


# ============================================================================
# SCRIPT ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    import sys
    
    # Get LaTeX file from command-line arguments or user input
    verbose_tag = False
    if len(sys.argv) > 2:
        if sys.argv[1] == '-v':
            verbose_tag = True
        latex_file = sys.argv[-1]
    elif len(sys.argv) > 1:
        if sys.argv[1] == '-v':
            verbose_tag = True
            latex_file = input("Enter the path to the LaTeX file: ").strip()
        else:
            latex_file = sys.argv[-1]
    else:
        latex_file = input("Enter the path to the LaTeX file: ").strip()
    
    # Resolve the path (handle both relative and absolute paths)
    latex_path = Path(latex_file)
    
    # If relative path, check current directory first, then parent directories
    if not latex_path.is_absolute():
        # Try current directory
        if not latex_path.exists():
            # Try common parent directories
            alt_paths = [
                Path.cwd() / latex_file,
                Path.cwd().parent / latex_file,
                Path.home() / latex_file,
            ]
            for alt_path in alt_paths:
                if alt_path.exists():
                    latex_path = alt_path
                    break
    
    # Check if file exists
    if not latex_path.exists():
        print(f"❌ Error: File '{latex_file}' not found")
        print(f"   Checked: {Path.cwd() / latex_file}")
        sys.exit(1)
    
    # Convert to absolute path for processing
    latex_file = str(latex_path.resolve())
    
    # Automatically generate output filename from LaTeX filename
    output_file = None  # This will use the default naming based on latex_file
    
    try:
        result = latex_to_word(latex_file, output_file, verbose=verbose_tag)
        if result:
            print(f"\n✅ Conversion completed successfully!")
            print(f"Output file: {result}")
        else:
            print("\n❌ Conversion failed")
    except Exception as e:
        print(f"\n❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
